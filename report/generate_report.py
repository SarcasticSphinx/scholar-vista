"""Generate the ScholarVista project report as a .docx file.

Run with the project's venv:
    .logo-venv/bin/python report/generate_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Theme ----------
PRIMARY = RGBColor(0x1E, 0x40, 0xAF)   # deep blue
ACCENT = RGBColor(0x64, 0x74, 0x8B)    # slate
LIGHT = RGBColor(0x94, 0xA3, 0xB8)

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def set_heading_color(paragraph, color=PRIMARY):
    for run in paragraph.runs:
        run.font.color.rgb = color


def add_heading(text, level=1, color=PRIMARY):
    h = doc.add_heading(text, level=level)
    set_heading_color(h, color)
    return h


def add_para(text, size=11, bold=False, italic=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_number(text):
    return doc.add_paragraph(text, style="List Number")


def add_image_placeholder(caption):
    """Insert a bordered box telling the user where to drop a screenshot."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    # shading
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F1F5F9")
    tcPr.append(shd)
    # borders
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "dashed")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), "94A3B8")
        borders.append(el)
    tcPr.append(borders)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n[  INSERT SCREENSHOT HERE  ]\n")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = ACCENT
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(caption + "\n")
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = LIGHT

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"Figure: {caption}")
    cr.italic = True
    cr.font.size = Pt(9)
    cr.font.color.rgb = ACCENT
    doc.add_paragraph()

# ============================================================
# TITLE / COVER PAGE
# ============================================================
for _ in range(4):
    doc.add_paragraph()

add_para("ScholarVista", size=34, bold=True, color=PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("A Full-Stack Scholarship Management Platform", size=16, italic=True,
         color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

add_para("Project Report", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

add_para("Submitted by: ____________________________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para("Student ID: ____________________________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para("Supervisor: ____________________________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para("Department: ____________________________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para("Date: ____________________________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (static)
# ============================================================
add_heading("Table of Contents", level=1)
toc_items = [
    "1. Introduction",
    "2. Problem Statement",
    "3. Project Objectives",
    "4. Scope of the Project",
    "5. System Users and Roles",
    "6. Key Features",
    "7. Technology Stack",
    "8. System Architecture",
    "9. Database Design",
    "10. Application Workflow",
    "11. Screenshots and Demonstration",
    "12. Testing and Quality Assurance",
    "13. Challenges Faced",
    "14. Future Enhancements",
    "15. Conclusion",
]
for item in toc_items:
    add_para(item, size=11, space_after=4)
doc.add_page_break()

# ============================================================
# 1. INTRODUCTION
# ============================================================
add_heading("1. Introduction", level=1)
add_para(
    "ScholarVista is a full-stack web application that centralizes the discovery, management, "
    "and application process for academic scholarships. The platform connects students who are "
    "seeking funding opportunities with universities and administrators who post and manage "
    "scholarship programs. It provides a single, searchable place where applicants can browse "
    "scholarships, compare options, bookmark opportunities, apply online, track application "
    "status, and leave reviews, while administrators and moderators manage listings, review "
    "submissions, and monitor platform activity."
)
add_para(
    "The system is built on a modern web stack using Next.js 16 with the App Router, a PostgreSQL "
    "database accessed through the Prisma ORM, and secure authentication provided by Better-Auth. "
    "The interface is styled with Tailwind CSS and Shadcn/UI components for a clean, responsive, "
    "and accessible user experience."
)

# ============================================================
# 2. PROBLEM STATEMENT
# ============================================================
add_heading("2. Problem Statement", level=1)
add_para(
    "Information about scholarships is typically scattered across many university websites, PDF "
    "notices, social media posts, and third-party blogs. This fragmentation creates several "
    "problems for students and institutions alike:"
)
add_bullet("Students struggle to find relevant scholarships and often miss deadlines because information is not centralized.", "Discovery difficulty: ")
add_bullet("It is hard to compare scholarships side by side in terms of coverage, stipend, eligibility, and fees.", "Lack of comparison: ")
add_bullet("Many scholarships require manual, offline, or email-based applications, making tracking difficult.", "Manual application process: ")
add_bullet("Applicants have no reliable way to track the status of their submissions or receive timely updates.", "No status visibility: ")
add_bullet("Universities and administrators lack a unified tool to publish, approve, and manage scholarship listings and applications.", "Fragmented administration: ")
add_para(
    "ScholarVista addresses these problems by providing a unified platform that brings scholarship "
    "discovery, comparison, online application, payment, status tracking, and administration "
    "together in one place."
)

# ============================================================
# 3. OBJECTIVES
# ============================================================
add_heading("3. Project Objectives", level=1)
add_number("To build a centralized platform where scholarships from multiple universities can be listed and searched.")
add_number("To allow students to browse, filter, compare, and bookmark scholarships easily.")
add_number("To enable students to apply online and track the real-time status of their applications.")
add_number("To provide a secure, role-based system for administrators and moderators to manage listings, approvals, and users.")
add_number("To support a review and rating system so applicants can share feedback about scholarships.")
add_number("To integrate payment handling for scholarships that require an application fee.")
add_number("To deliver notifications that keep users informed of status changes and important events.")

# ============================================================
# 4. SCOPE
# ============================================================
add_heading("4. Scope of the Project", level=1)
add_para("The scope of ScholarVista covers the following areas:")
add_bullet("Public browsing of scholarships and universities without requiring an account.")
add_bullet("User registration, authentication, and profile management.")
add_bullet("Scholarship application, bookmarking, comparison, and review functionality for registered users.")
add_bullet("A dashboard for administrators and moderators to manage scholarships, universities, users, applications, approvals, reports, and platform settings.")
add_bullet("Payment processing for scholarships that charge an application fee, including handling of expired payments.")
add_bullet("Notification delivery for application status changes, scholarship approvals, and payment confirmations.")

# ============================================================
# 5. USERS AND ROLES
# ============================================================
add_heading("5. System Users and Roles", level=1)
add_para("ScholarVista defines three distinct roles, each with different permissions:")

table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].paragraphs[0].add_run("Role").bold = True
hdr[1].paragraphs[0].add_run("Responsibilities").bold = True
roles = [
    ("User (Student)", "Browses and compares scholarships, bookmarks opportunities, submits applications, makes payments, tracks application status, and writes reviews."),
    ("Moderator", "Reviews and approves scholarship submissions, manages listings and applications, and helps maintain content quality."),
    ("Admin", "Has full control of the platform, including managing users, universities, scholarships, approvals, reports, and platform-wide settings."),
]
for role, desc in roles:
    row = table.add_row().cells
    row[0].paragraphs[0].add_run(role).bold = True
    row[1].text = desc
doc.add_paragraph()

# ============================================================
# 6. KEY FEATURES
# ============================================================
add_heading("6. Key Features", level=1)

add_heading("6.1 Public and Student Features", level=2)
add_bullet("Browse and search a catalog of scholarships and partner universities.")
add_bullet("Filter scholarships by category (Undergraduate, Masters, PhD, Postdoc, Exchange, Short Course) and other criteria.")
add_bullet("Compare multiple scholarships side by side.")
add_bullet("Bookmark scholarships to revisit later.")
add_bullet("Submit online applications with academic and personal details.")
add_bullet("Track application status (Pending, Processing, Completed, Rejected).")
add_bullet("Write and read reviews with star ratings for scholarships.")
add_bullet("Receive in-app notifications for important events.")
add_bullet("Manage a personal profile including educational level, major, and location.")

add_heading("6.2 Administration and Moderation Features", level=2)
add_bullet("Role-based dashboard for admins and moderators.")
add_bullet("Create, edit, and manage scholarship listings and universities.")
add_bullet("Approve or reject user-submitted scholarships.")
add_bullet("Manage user accounts and roles.")
add_bullet("Review and update the status of applications.")
add_bullet("View reports and analytics on platform activity.")
add_bullet("Configure platform-wide settings such as enabling user submissions and payments.")

add_heading("6.3 Payments and Notifications", level=2)
add_bullet("Handle application fees through a checkout and payment flow.")
add_bullet("Track payment status (Unpaid, Paid, Refunded, Failed, Expired).")
add_bullet("Automatically expire pending payments via a scheduled job.")
add_bullet("Send notifications on application status changes, scholarship approvals, and payment confirmations.")

# ============================================================
# 7. TECHNOLOGY STACK
# ============================================================
add_heading("7. Technology Stack", level=1)
tech_table = doc.add_table(rows=1, cols=2)
tech_table.style = "Light Grid Accent 1"
th = tech_table.rows[0].cells
th[0].paragraphs[0].add_run("Layer / Concern").bold = True
th[1].paragraphs[0].add_run("Technology").bold = True
tech = [
    ("Framework", "Next.js 16 (App Router) with React 19"),
    ("Language", "TypeScript"),
    ("Styling", "Tailwind CSS v4 with Shadcn/UI (Radix UI) components"),
    ("Authentication", "Better-Auth (secure, type-safe)"),
    ("Database", "PostgreSQL"),
    ("ORM", "Prisma"),
    ("Validation", "Zod schema validation"),
    ("File Uploads", "UploadThing"),
    ("Icons", "Lucide React"),
    ("Charts", "Recharts (reports and analytics)"),
    ("Testing", "Vitest, Testing Library, Playwright (E2E), Axe (accessibility)"),
    ("Tooling", "ESLint, Prettier"),
]
for layer, technology in tech:
    row = tech_table.add_row().cells
    row[0].paragraphs[0].add_run(layer).bold = True
    row[1].text = technology
doc.add_paragraph()

# ============================================================
# 8. ARCHITECTURE
# ============================================================
add_heading("8. System Architecture", level=1)
add_para(
    "ScholarVista follows the Next.js App Router architecture, combining server-rendered pages, "
    "React Server Components, and Server Actions. The application is organized into route groups "
    "that separate concerns:"
)
add_bullet("Public pages for unauthenticated browsing of scholarships and universities.", "Public area: ")
add_bullet("Authenticated pages for logged-in students (applications, bookmarks, reviews, notifications, profile).", "Authenticated area: ")
add_bullet("Role-protected dashboard for admins and moderators.", "Dashboard area: ")
add_bullet("Server Actions and API routes handle data mutations, authentication, uploads, webhooks, and scheduled jobs.", "Backend logic: ")
add_para(
    "Business logic is implemented primarily through Server Actions (in the actions layer), which "
    "interact with the database through Prisma. Authentication and route protection are enforced "
    "through Better-Auth and middleware. Input is validated using Zod before it reaches the "
    "database."
)

# ============================================================
# 9. DATABASE DESIGN
# ============================================================
add_heading("9. Database Design", level=1)
add_para(
    "The data model is managed with Prisma and stored in PostgreSQL. The core entities and their "
    "relationships are summarized below."
)
db_table = doc.add_table(rows=1, cols=2)
db_table.style = "Light Grid Accent 1"
dh = db_table.rows[0].cells
dh[0].paragraphs[0].add_run("Entity").bold = True
dh[1].paragraphs[0].add_run("Description").bold = True
entities = [
    ("User", "Represents students, moderators, and admins. Stores profile details and role."),
    ("University", "Institutions offering scholarships, including type, rank, and partner status."),
    ("Scholarship", "Scholarship listings with category, coverage, stipend, deadline, fees, and approval status."),
    ("Application", "A student's application to a scholarship, with academic details, status, and payment status."),
    ("Review", "A star rating (1-5) and comment left by a user for a scholarship."),
    ("Payment", "Records payment transactions for application fees, with status and expiry."),
    ("Notification", "In-app messages sent to users about status changes and events."),
    ("Bookmark", "Links a user to a saved scholarship."),
    ("PlatformSettings", "Global configuration such as toggles for user submissions and payments."),
]
for ent, desc in entities:
    row = db_table.add_row().cells
    row[0].paragraphs[0].add_run(ent).bold = True
    row[1].text = desc
doc.add_paragraph()
add_para("Key relationships:", bold=True)
add_bullet("A University has many Scholarships; a Scholarship belongs to one University.")
add_bullet("A User can post many Scholarships and submit many Applications.")
add_bullet("A Scholarship has many Applications, Reviews, Bookmarks, and Payments.")
add_bullet("Each user can apply, bookmark, or review a given scholarship only once (enforced by unique constraints).")

add_image_placeholder("Entity-Relationship Diagram (ERD) of the database")

# ============================================================
# 10. WORKFLOW
# ============================================================
add_heading("10. Application Workflow", level=1)
add_para("A typical end-to-end flow for a student on the platform is as follows:")
add_number("The student registers or signs in to the platform.")
add_number("They browse or search scholarships and can compare or bookmark options.")
add_number("They open a scholarship and submit an online application with their details.")
add_number("If the scholarship requires a fee, they complete the payment through the checkout flow.")
add_number("The application enters the review pipeline (Pending to Processing).")
add_number("A moderator or admin reviews the application and updates its status.")
add_number("The student receives a notification about the status change and can track it from their dashboard.")

# ============================================================
# 11. SCREENSHOTS
# ============================================================
add_heading("11. Screenshots and Demonstration", level=1)
add_para(
    "This section demonstrates the main pages of the application. Insert the corresponding "
    "screenshot into each placeholder below."
)

shots = [
    ("Home / Landing page", "The public landing page introducing ScholarVista."),
    ("Scholarship listing page", "The browsable catalog of scholarships with filters and search."),
    ("Scholarship details page", "Detailed view of a single scholarship, including reviews."),
    ("Scholarship comparison page", "Side-by-side comparison of selected scholarships."),
    ("Universities page", "The list of partner universities."),
    ("Sign-in / Sign-up page", "The authentication screens powered by Better-Auth."),
    ("Application form", "The online scholarship application form."),
    ("Checkout / Payment page", "The payment flow for application fees."),
    ("My Applications page", "A student's applications with their current status."),
    ("My Bookmarks page", "Scholarships the student has saved."),
    ("Notifications page", "In-app notifications for the user."),
    ("User profile page", "Editable student profile details."),
    ("Admin / Moderator dashboard", "Overview dashboard with key metrics."),
    ("Manage scholarships (dashboard)", "Admin view for creating and editing scholarships."),
    ("Approvals page", "Approving or rejecting user-submitted scholarships."),
    ("Reports and analytics page", "Charts summarizing platform activity."),
]
for i, (title, caption) in enumerate(shots, start=1):
    add_heading(f"11.{i} {title}", level=2)
    add_image_placeholder(caption)

# ============================================================
# 12. TESTING
# ============================================================
add_heading("12. Testing and Quality Assurance", level=1)
add_para("The project includes multiple layers of testing and quality tooling:")
add_bullet("Unit and component tests using Vitest and React Testing Library.", "Unit tests: ")
add_bullet("Integration tests using a dedicated Vitest configuration.", "Integration tests: ")
add_bullet("End-to-end tests using Playwright to validate full user flows.", "E2E tests: ")
add_bullet("Automated accessibility checks using Axe with Playwright.", "Accessibility: ")
add_bullet("ESLint and TypeScript type-checking to enforce code quality and type safety.", "Static analysis: ")

# ============================================================
# 13. CHALLENGES
# ============================================================
add_heading("13. Challenges Faced", level=1)
add_bullet("Designing a normalized database schema that supports scholarships, applications, payments, and reviews while enforcing constraints such as one application per user per scholarship.")
add_bullet("Implementing secure, role-based access control across public, authenticated, and dashboard areas.")
add_bullet("Handling the payment lifecycle, including expiring pending payments through a scheduled job.")
add_bullet("Keeping the interface responsive and accessible across devices.")
add_bullet("Managing complex state and data flows using Server Actions and Server Components in Next.js 16.")

# ============================================================
# 14. FUTURE ENHANCEMENTS
# ============================================================
add_heading("14. Future Enhancements", level=1)
add_bullet("Email and push notifications in addition to in-app notifications.")
add_bullet("AI-based scholarship recommendations tailored to a student's profile.")
add_bullet("Advanced analytics and exportable reports for administrators.")
add_bullet("Multi-language support for a wider audience.")
add_bullet("Integration with additional payment gateways.")
add_bullet("A mobile application for on-the-go access.")

# ============================================================
# 15. CONCLUSION
# ============================================================
add_heading("15. Conclusion", level=1)
add_para(
    "ScholarVista successfully delivers a unified platform for scholarship discovery and "
    "management. It solves the problem of scattered scholarship information by centralizing "
    "listings, streamlining the application process, and providing tools for both students and "
    "administrators. Built with a modern, type-safe technology stack and supported by a "
    "comprehensive testing setup, the project demonstrates a complete full-stack application "
    "that is both practical and extensible. The planned future enhancements provide a clear path "
    "for continued growth beyond the current scope."
)

# ---------- Save ----------
import os
out_path = os.path.join(os.path.dirname(__file__), "ScholarVista_Project_Report.docx")
doc.save(out_path)
print("Saved:", out_path)
